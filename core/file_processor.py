import os
import zipfile
import tempfile
import shutil
import re
import time
import logging
from typing import List, Dict, Tuple, Any

# Set up logging
logger = logging.getLogger('file_processing')


class FileProcessor:
    """Process myITS Classroom ZIP files and extract student submissions."""

    def __init__(self, temp_dir: str = None):
        """Initialize the file processor.

        Args:
            temp_dir: Directory to store extracted files temporarily
        """
        self.temp_dir = temp_dir or tempfile.mkdtemp()
        logger.info(f"FileProcessor initialized with temp_dir: {self.temp_dir}")

        # Enhanced file type mappings
        self.language_extensions = {
            'python': ['.py', '.pyw'],
            'java': ['.java'],
            'c': ['.c', '.h'],
            'cpp': ['.cpp', '.cc', '.cxx', '.c++', '.hpp', '.hh', '.hxx', '.h++'],
            'javascript': ['.js', '.jsx', '.mjs'],
            'csharp': ['.cs'],
            'php': ['.php'],
            'ruby': ['.rb'],
            'go': ['.go'],
            'rust': ['.rs'],
            'swift': ['.swift'],
            'kotlin': ['.kt', '.kts'],
            'typescript': ['.ts', '.tsx'],
            'r': ['.r', '.R'],
            'matlab': ['.m'],
            'scala': ['.scala'],
            'perl': ['.pl', '.pm'],
            'shell': ['.sh', '.bash', '.zsh', '.fish'],
            'sql': ['.sql'],
            'html': ['.html', '.htm', '.xhtml'],
            'css': ['.css', '.scss', '.sass', '.less'],
            'xml': ['.xml', '.xsl', '.xsd'],
            'json': ['.json'],
            'yaml': ['.yml', '.yaml'],
            'markdown': ['.md', '.markdown']
        }

        self.text_extensions = [
            '.txt', '.rtf', '.pdf', '.doc', '.docx', '.odt',
            '.md', '.markdown', '.tex', '.latex'
        ]

        self.archive_extensions = ['.zip', '.tar', '.gz', '.bz2', '.7z', '.rar']

        self.image_extensions = [
            '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff',
            '.svg', '.webp', '.ico'
        ]

        logger.debug(f"Configured {len(self.language_extensions)} language mappings")
        logger.debug(f"Text extensions: {len(self.text_extensions)}")
        logger.debug(f"Archive extensions: {len(self.archive_extensions)}")
        logger.debug(f"Image extensions: {len(self.image_extensions)}")

    def extract_zip(self, zip_path: str) -> str:
        """Extract the ZIP file to a temporary directory.

        Args:
            zip_path: Path to the ZIP file

        Returns:
            Path to the extracted directory
        """
        start_time = time.time()
        logger.info(f"Starting ZIP extraction: {zip_path}")

        # Check if file exists and get size
        if not os.path.exists(zip_path):
            error_msg = f"ZIP file not found: {zip_path}"
            logger.error(error_msg)
            raise FileNotFoundError(error_msg)

        file_size = os.path.getsize(zip_path)
        logger.info(f"ZIP file size: {file_size / 1024 / 1024:.2f} MB")

        extract_dir = os.path.join(self.temp_dir, 'extracted')
        os.makedirs(extract_dir, exist_ok=True)
        logger.debug(f"Created extraction directory: {extract_dir}")

        try:
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                # Log ZIP contents summary
                file_list = zip_ref.namelist()
                logger.info(f"ZIP contains {len(file_list)} files")

                # Count file types
                extensions = {}
                for file_name in file_list:
                    ext = os.path.splitext(file_name)[1].lower()
                    extensions[ext] = extensions.get(ext, 0) + 1

                logger.debug(f"File extensions in ZIP: {extensions}")

                # Extract all files
                zip_ref.extractall(extract_dir)

            extraction_time = time.time() - start_time
            logger.info(f"ZIP extraction completed in {extraction_time:.2f} seconds")
            logger.info(f"Extracted to: {extract_dir}")

            return extract_dir

        except zipfile.BadZipFile as e:
            error_msg = f"Invalid ZIP file: {e}"
            logger.error(error_msg)
            raise
        except Exception as e:
            error_msg = f"ZIP extraction error: {e}"
            logger.error(error_msg)
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            raise

    def parse_moodle_structure(self, extract_dir: str) -> List[Dict[str, Any]]:
        """Parse the Moodle directory structure to extract student submissions.

        Enhanced version with better file type detection and language identification.

        Args:
            extract_dir: Path to the extracted directory

        Returns:
            List of dictionaries containing student info and submission paths
        """
        start_time = time.time()
        logger.info(f"Starting Moodle structure parsing: {extract_dir}")

        submissions = []

        # Enhanced regex patterns for Moodle directory structures
        file_patterns = [
            r'(.+)_(\d+)_assignsubmission_file_',
            r'(.+)_(\d+)_file_',
            r'(.+?)_(\d+)_'
        ]

        onlinetext_patterns = [
            r'(.+)_(\d+)_assignsubmission_onlinetext_',
            r'(.+)_(\d+)_onlinetext_'
        ]

        logger.debug(f"Using file patterns: {file_patterns}")
        logger.debug(f"Using onlinetext patterns: {onlinetext_patterns}")

        # Walk through directory structure
        total_dirs = 0
        processed_dirs = 0

        for root, dirs, files in os.walk(extract_dir):
            total_dirs += len(dirs)

            # Process file submissions
            for dir_name in dirs:
                logger.debug(f"Processing directory: {dir_name}")
                processed_dirs += 1

                student_name = None
                student_id = None

                # Try different file patterns
                for i, pattern in enumerate(file_patterns):
                    match = re.match(pattern, dir_name)
                    if match:
                        student_name = match.group(1)
                        student_id = match.group(2)
                        logger.debug(f"Matched file pattern {i + 1}: {student_name} ({student_id})")
                        break

                if student_name and student_id:
                    submission_dir = os.path.join(root, dir_name)
                    submission_files = []

                    logger.debug(f"Processing submission for {student_name}")

                    # Get all files in the submission directory
                    file_count = 0
                    for sub_root, _, sub_files in os.walk(submission_dir):
                        for file in sub_files:
                            file_count += 1
                            file_path = os.path.join(sub_root, file)
                            file_ext = os.path.splitext(file)[1].lower()

                            logger.debug(f"  Analyzing file: {file} ({file_ext})")

                            # Enhanced file analysis
                            file_info = self._analyze_file(file, file_path, file_ext)
                            submission_files.append(file_info)

                            logger.debug(
                                f"    Type: {file_info['type']}, Language: {file_info['language']}, Size: {file_info['size']} bytes")

                    if submission_files:
                        cleaned_name = self._clean_student_name(student_name)
                        logger.info(f"Found submission: {cleaned_name} ({student_id}) - {len(submission_files)} files")

                        submissions.append({
                            'student_name': cleaned_name,
                            'student_id': student_id,
                            'submission_dir': submission_dir,
                            'files': submission_files,
                            'submission_type': 'file'
                        })
                    else:
                        logger.warning(f"No files found for {student_name}")

                # Process online text submissions
                for i, pattern in enumerate(onlinetext_patterns):
                    onlinetext_match = re.match(pattern, dir_name)
                    if onlinetext_match:
                        student_name = onlinetext_match.group(1)
                        student_id = onlinetext_match.group(2)
                        logger.debug(f"Matched onlinetext pattern {i + 1}: {student_name} ({student_id})")

                        submission_dir = os.path.join(root, dir_name)
                        online_text_files = []

                        # Look for onlinetext.html files
                        for sub_root, _, sub_files in os.walk(submission_dir):
                            for file in sub_files:
                                if file.lower() in ['onlinetext.html', 'onlinetext.htm']:
                                    file_path = os.path.join(sub_root, file)
                                    logger.debug(f"Found online text file: {file_path}")

                                    try:
                                        text_content = self.extract_text_from_html(file_path)
                                        logger.debug(f"Extracted {len(text_content)} characters from HTML")

                                        online_text_files.append({
                                            'name': file,
                                            'path': file_path,
                                            'extension': '.html',
                                            'content': text_content,
                                            'type': 'text',
                                            'language': 'html',
                                            'size': os.path.getsize(file_path),
                                            'is_code': False,
                                            'is_text': True
                                        })
                                    except Exception as e:
                                        logger.error(f"Error extracting text from {file_path}: {e}")

                        if online_text_files:
                            cleaned_name = self._clean_student_name(student_name)
                            logger.info(
                                f"Found online text submission: {cleaned_name} ({student_id}) - {len(online_text_files)} files")

                            submissions.append({
                                'student_name': cleaned_name,
                                'student_id': student_id,
                                'submission_dir': submission_dir,
                                'files': online_text_files,
                                'submission_type': 'onlinetext'
                            })

        parsing_time = time.time() - start_time
        logger.info(f"Moodle structure parsing completed in {parsing_time:.2f} seconds")
        logger.info(f"Processed {processed_dirs} directories, found {len(submissions)} valid submissions")

        # Log submission summary
        submission_types = {}
        total_files = 0
        for submission in submissions:
            sub_type = submission['submission_type']
            submission_types[sub_type] = submission_types.get(sub_type, 0) + 1
            total_files += len(submission['files'])

        logger.info(f"Submission types: {submission_types}")
        logger.info(f"Total files processed: {total_files}")

        return submissions

    def _clean_student_name(self, name: str) -> str:
        """Clean and normalize student names."""
        # Remove common Moodle artifacts
        name = re.sub(r'_\d+$', '', name)  # Remove trailing numbers
        name = name.replace('_', ' ')  # Replace underscores with spaces
        name = ' '.join(word.capitalize() for word in name.split())  # Proper case
        return name.strip()

    def _analyze_file(self, filename: str, filepath: str, extension: str) -> Dict[str, Any]:
        """Analyze a file and determine its properties."""
        file_size = 0
        try:
            file_size = os.path.getsize(filepath)
        except:
            pass

        # Determine file language and type
        language = self._get_file_language(extension)
        file_type = self._determine_enhanced_file_type(extension, language)
        is_code = language in self.language_extensions
        is_text = extension in self.text_extensions or language in ['html', 'css', 'xml', 'json', 'yaml', 'markdown']

        # Additional analysis for code files
        complexity_score = 0
        if is_code:
            complexity_score = self._estimate_code_complexity(filepath)

        return {
            'name': filename,
            'path': filepath,
            'extension': extension,
            'type': file_type,
            'language': language,
            'size': file_size,
            'is_code': is_code,
            'is_text': is_text,
            'complexity_score': complexity_score
        }

    def _get_file_language(self, extension: str) -> str:
        """Determine programming language from file extension."""
        extension = extension.lower()

        for language, extensions in self.language_extensions.items():
            if extension in extensions:
                return language

        # Special cases
        if extension in self.text_extensions:
            return 'text'
        elif extension in self.image_extensions:
            return 'image'
        elif extension in self.archive_extensions:
            return 'archive'
        else:
            return 'unknown'

    def _determine_enhanced_file_type(self, extension: str, language: str) -> str:
        """Enhanced file type determination."""
        if language in ['python', 'java', 'c', 'cpp', 'javascript', 'csharp', 'php', 'ruby', 'go', 'rust', 'swift',
                        'kotlin']:
            return 'code'
        elif language in ['html', 'css', 'xml', 'json', 'yaml']:
            return 'markup'
        elif language in ['markdown', 'text'] or extension in self.text_extensions:
            return 'text'
        elif language == 'sql':
            return 'database'
        elif language in ['shell']:
            return 'script'
        elif language == 'image':
            return 'image'
        elif language == 'archive':
            return 'archive'
        else:
            return 'other'

    def _estimate_code_complexity(self, filepath: str) -> int:
        """Estimate code complexity based on file content."""
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            # Simple complexity metrics
            lines = content.split('\n')
            non_empty_lines = [line.strip() for line in lines if line.strip()]

            # Count control structures
            control_keywords = ['if', 'else', 'elif', 'for', 'while', 'switch', 'case', 'try', 'catch', 'finally']
            control_count = sum(content.lower().count(keyword) for keyword in control_keywords)

            # Function/method count
            function_patterns = [r'def\s+\w+', r'function\s+\w+', r'\w+\s*\([^)]*\)\s*{', r'public\s+\w+\s+\w+\s*\(']
            function_count = sum(len(re.findall(pattern, content, re.IGNORECASE)) for pattern in function_patterns)

            # Simple complexity score
            complexity = len(non_empty_lines) + (control_count * 2) + (function_count * 3)
            return min(complexity, 100)  # Cap at 100

        except:
            return 0

    def _determine_file_type(self, extension: str) -> str:
        """Legacy method for backward compatibility."""
        language = self._get_file_language(extension)
        return self._determine_enhanced_file_type(extension, language)

    def get_file_content(self, file_path: str) -> str:
        """Enhanced method to get the content of a file."""
        logger.debug(f"Reading file content: {file_path}")

        if not os.path.exists(file_path):
            error_msg = f"File not found: {file_path}"
            logger.error(error_msg)
            return f"Error: {error_msg}"

        file_size = os.path.getsize(file_path)
        logger.debug(f"File size: {file_size} bytes")

        ext = os.path.splitext(file_path)[1].lower()
        logger.debug(f"File extension: {ext}")

        # Handle code files directly
        if ext in ['.py', '.java', '.c', '.cpp', '.cc', '.cxx', '.c++', '.js', '.jsx', '.ts', '.tsx',
                   '.cs', '.php', '.rb', '.go', '.rs', '.swift', '.kt', '.html', '.htm', '.css',
                   '.xml', '.json', '.yml', '.yaml', '.md', '.sh', '.sql', '.r', '.m', '.scala', '.pl']:
            logger.debug(f"Processing as code/text file: {ext}")
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                logger.debug(f"Successfully read {len(content)} characters")
                logger.debug(f"Content preview: {content[:100]}...")
                return content
            except UnicodeDecodeError:
                logger.warning(f"UTF-8 decode failed, trying latin-1: {file_path}")
                try:
                    with open(file_path, 'r', encoding='latin-1') as f:
                        content = f.read()
                    logger.debug(f"Successfully read with latin-1: {len(content)} characters")
                    return content
                except Exception as e:
                    error_msg = f"Could not read file {file_path}: {e}"
                    logger.error(error_msg)
                    return f"Error: {error_msg}"
            except Exception as e:
                error_msg = f"Error reading file {file_path}: {e}"
                logger.error(error_msg)
                return f"Error: {error_msg}"

        # Handle text files
        elif ext in ['.txt', '.md']:
            logger.debug(f"Processing as plain text file: {ext}")
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                logger.debug(f"Successfully read text file: {len(content)} characters")
                return content
            except Exception as e:
                error_msg = f"Could not read text file {file_path}: {e}"
                logger.error(error_msg)
                return f"Error: {error_msg}"

        # Handle document formats
        else:
            logger.debug(f"Processing as document file: {ext}")
            return self.extract_text_from_file(file_path)

    def extract_text_from_html(self, file_path: str) -> str:
        """Extract text content from HTML file with improved handling."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()

            # Try BeautifulSoup first
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html_content, 'html.parser')

                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()

                # Get text and clean it up
                text = soup.get_text(separator='\n')
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)

                return text

            except ImportError:
                # Fallback regex method
                import re
                # Remove script and style elements
                html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)

                # Remove HTML tags
                text = re.sub(r'<[^>]+>', ' ', html_content)

                # Clean up whitespace
                text = re.sub(r'\s+', ' ', text)
                return text.strip()

        except Exception as e:
            return f"Error extracting HTML content: {str(e)}"

    def extract_text_from_file(self, file_path: str) -> str:
        """Enhanced text extraction from various file formats."""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.txt':
            try:
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            return f.read()
                    except UnicodeDecodeError:
                        continue
                return "Error: Could not decode text file"
            except Exception as e:
                return f"Error reading text file: {str(e)}"

        elif ext == '.docx':
            try:
                import docx
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)

                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            full_text.append(cell.text)

                return "\n".join(full_text)
            except ImportError:
                return "Error: python-docx package required for DOCX processing"
            except Exception as e:
                return f"Error processing DOCX file: {str(e)}"

        elif ext == '.pdf':
            try:
                # Try PyPDF2 first
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() + "\n"
                        return text
                except:
                    # Fallback to pdfplumber if available
                    try:
                        import pdfplumber
                        with pdfplumber.open(file_path) as pdf:
                            text = ""
                            for page in pdf.pages:
                                text += page.extract_text() + "\n"
                            return text
                    except ImportError:
                        return "Error: PyPDF2 or pdfplumber package required for PDF processing"
            except Exception as e:
                return f"Error processing PDF file: {str(e)}"

        elif ext == '.odt':
            try:
                from odf.opendocument import load
                from odf.text import P
                doc = load(file_path)
                paras = doc.getElementsByType(P)
                text_content = []
                for para in paras:
                    para_text = str(para)
                    # Simple text extraction from ODT
                    import re
                    clean_text = re.sub(r'<[^>]+>', '', para_text)
                    if clean_text.strip():
                        text_content.append(clean_text.strip())
                return "\n".join(text_content)
            except ImportError:
                return "Error: odfpy package required for ODT processing"
            except Exception as e:
                return f"Error processing ODT file: {str(e)}"

        elif ext == '.rtf':
            try:
                from striprtf.striprtf import rtf_to_text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    rtf_content = f.read()
                return rtf_to_text(rtf_content)
            except ImportError:
                return "Error: striprtf package required for RTF processing"
            except Exception as e:
                return f"Error processing RTF file: {str(e)}"

        else:
            return f"Unsupported file format: {ext}"

    def categorize_submissions(self, submissions: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
        """Enhanced categorization of submissions by file type and language."""
        categorized = {
            'code': [],
            'text': [],
            'markup': [],
            'script': [],
            'database': [],
            'other': []
        }

        # Also categorize by programming language
        language_categorized = {}

        for submission in submissions:
            # Handle online text submissions
            if submission['submission_type'] == 'onlinetext':
                for file in submission['files']:
                    categorized['text'].append({**submission, 'current_file': file})
                continue

            # Handle file submissions
            for file in submission['files']:
                file_type = file.get('type', 'other')
                language = file.get('language', 'unknown')

                # Add to main categories
                if file_type in categorized:
                    categorized[file_type].append({**submission, 'current_file': file})
                else:
                    categorized['other'].append({**submission, 'current_file': file})

                # Add to language-specific categories
                if language != 'unknown' and file.get('is_code', False):
                    if language not in language_categorized:
                        language_categorized[language] = []
                    language_categorized[language].append({**submission, 'current_file': file})

        # Add language categories to the result
        categorized.update(language_categorized)

        return categorized

    def get_submission_statistics(self, submissions: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Get detailed statistics about submissions."""
        stats = {
            'total_submissions': len(submissions),
            'total_files': 0,
            'file_types': {},
            'languages': {},
            'average_files_per_student': 0,
            'largest_file': {'name': '', 'size': 0, 'student': ''},
            'most_complex_code': {'name': '', 'complexity': 0, 'student': ''},
            'submission_types': {'file': 0, 'onlinetext': 0}
        }

        total_files = 0
        total_complexity = 0
        code_files = 0

        for submission in submissions:
            stats['submission_types'][submission.get('submission_type', 'file')] += 1

            for file_info in submission['files']:
                total_files += 1

                # File type statistics
                file_type = file_info.get('type', 'other')
                stats['file_types'][file_type] = stats['file_types'].get(file_type, 0) + 1

                # Language statistics
                language = file_info.get('language', 'unknown')
                if language != 'unknown':
                    stats['languages'][language] = stats['languages'].get(language, 0) + 1

                # Size tracking
                file_size = file_info.get('size', 0)
                if file_size > stats['largest_file']['size']:
                    stats['largest_file'] = {
                        'name': file_info['name'],
                        'size': file_size,
                        'student': submission['student_name']
                    }

                # Complexity tracking for code files
                if file_info.get('is_code', False):
                    code_files += 1
                    complexity = file_info.get('complexity_score', 0)
                    total_complexity += complexity

                    if complexity > stats['most_complex_code']['complexity']:
                        stats['most_complex_code'] = {
                            'name': file_info['name'],
                            'complexity': complexity,
                            'student': submission['student_name']
                        }

        stats['total_files'] = total_files
        stats['average_files_per_student'] = total_files / len(submissions) if submissions else 0
        stats['average_code_complexity'] = total_complexity / code_files if code_files > 0 else 0

        return stats

    def cleanup(self):
        """Clean up temporary files."""
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)